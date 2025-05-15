from pypdf import PdfReader
import os
from docx import Document
import openpyxl
import msoffcrypto
from typing import Optional
from .models import ExtractionResult
import io
from msoffcrypto import OfficeFile
import tempfile  

def parse_document(file_path: str, password: Optional[str] = None) -> ExtractionResult:
    if file_path.endswith('.pdf'):
        return parse_pdf(file_path, password)
    elif file_path.endswith('.docx'):
        return parse_docx(file_path, password)  
    elif file_path.endswith(('.xlsx', '.xls')):
        return parse_excel(file_path, password)  
    else:
        raise ValueError("Unsupported file format")

def parse_pdf(file_path: str, password: Optional[str]) -> ExtractionResult:
    content = ""
    metadata = {"type": "pdf", "pages": 0, "encrypted": False}
    
    try:
        reader = PdfReader(file_path)
        if reader.is_encrypted:
            metadata["encrypted"] = True
            if not password:
                raise ValueError("Password required for PDF")
            if not reader.decrypt(password):
                raise ValueError("Incorrect PDF password")
        
        metadata["pages"] = len(reader.pages)
        for page in reader.pages:
            content += page.extract_text() + "\n\n"
            
    except Exception as e:
        raise ValueError(f"PDF processing error: {str(e)}")
    
    return ExtractionResult(content=content.strip(), metadata=metadata)

def parse_docx(file_path: str, password: Optional[str] = None) -> ExtractionResult:
    content = ""
    metadata = {"type": "docx", "encrypted": False}
    
    try:
        with open(file_path, "rb") as f:
            if password:
                metadata["encrypted"] = True
                decrypted = io.BytesIO()
                office_file = OfficeFile(f)
                office_file.load_key(password=password)
                office_file.decrypt(decrypted)
                doc = Document(decrypted)
            else:
                doc = Document(f)
        
        content = "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        raise ValueError(f"DOCX processing error: {str(e)}")
    
    return ExtractionResult(content=content.strip(), metadata=metadata)

def parse_excel(file_path: str, password: Optional[str] = None) -> ExtractionResult:
    content = []
    metadata = {"type": "excel", "sheets": 0, "encrypted": False}
    
    try:
        if password:
            metadata["encrypted"] = True
            temp_decrypted = tempfile.NamedTemporaryFile(delete=False)
            try:
                with open(file_path, "rb") as f:
                    office_file = OfficeFile(f)
                    office_file.load_key(password=password)
                    office_file.decrypt(temp_decrypted)
                
                wb = openpyxl.load_workbook(temp_decrypted.name)
            finally:
                temp_decrypted.close()
                os.unlink(temp_decrypted.name)
        else:
            wb = openpyxl.load_workbook(file_path)
        
        metadata["sheets"] = len(wb.sheetnames)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            content.append(f"\n--- Sheet: {sheet_name} ---\n")
            for row in sheet.iter_rows(values_only=True):
                row_content = " | ".join(str(cell) if cell is not None else "" for cell in row)
                content.append(row_content)
                
    except Exception as e:
        raise ValueError(f"Excel processing error: {str(e)}")
    
    return ExtractionResult(content="\n".join(content), metadata=metadata)